import random
import pandas
from statistics import mean, median, stdev
from datetime import datetime, date, timedelta
from docx import Document


List_MD = []
List_Services = []
statistic_list = []


class ChurchServers:
    def __init__(self, lastname, firstname, abbreviation, unavailable=[], group="Null"):
        # MD is used as an abbreviation of ChurchServers throughout the code
        # Different Name Attribute of the church server
        self.lastname = lastname
        self.firstname = firstname
        self.fullname = f"{self.firstname} {self.lastname}"
        self.abbreviation = abbreviation
        # Shows if someone is already in the current Messe
        self.is_allocated = True
        # Shows if someone has excused themselves
        self.is_available = False
        # Shows if someone is able to serve at Church
        self.is_advanced = True
        # Groups are used to organize the ChurchServers and make them more manageable
        # TODO Selecting two Church Servers in the same group will be preferred
        # "Null" is a String to show the group did not work
        self.group = str(group)
        self.counter = 0
        # List of TimeSpan objects when a ChurchServer is not available
        self.unavailable = unavailable

    def add_unavailable(self, timespan):
        overlap_list = []
        for selected_unavailable in range(len(self.unavailable)):
            if timespan.check_overlap(self.unavailable[selected_unavailable]):  # Checks if there are any overlaps
                overlap_list.append(self.unavailable[selected_unavailable]) # Mergeable TimeSpans are appended to a list

        for selected_merge in range(len(overlap_list)):
            self.unavailable.remove(overlap_list[selected_merge]) # removes the overlapping element
            timespan = merge_timespan(timespan, overlap_list[selected_merge]) # merges all overlapping elements
        self.unavailable.append(timespan)  # adds the final TimeSpan to the unavailable list



class ChurchService:
    def __init__(self, number_md_needed, date, time = "Null"):
        # numberMD is the number of ChurchServers needed for the Church Service
        self.count = number_md_needed
        self.ListServingMD = []
        # NumberAllocatedMD shows the current amount of Church Servers that are allocated to the Service.
        # If the allocation is finished this number should be equal to number_md_needed
        self.NumberAllocatedMD = len(self.ListServingMD)
        # Adds a time to each ChurchService
        # date_time should be a datetime object
        #TODO readd the datetime of a Service
        self.time = time
        # self.date is a Datetime object as a .isoformat() string
        self.date = date
        self.description = f"Messe am {self.date} mit {self.count} Messdienern"


class TimeSpan:
    def __init__(self, start_date, end_date):
        self.start_date = start_date
        self.end_date = end_date
        self.order_dates()
        self.description = f"{self.start_date} / {self.end_date}"

    def during_timespan(self, checked_date): # Takes a date and returns true if it is in the TimeSpan
        if self.start_date <= checked_date <= self.end_date:
            return True
        else:
            return False

    def check_overlap(self, other):
        # TODO add functionality if day is one of start or end date
        if other.during_timespan(self.start_date - timedelta(days=1)):
            return True
        elif self.during_timespan(other.start_date):
            # This is needed because otherwise one can add a one day TimeSpan to itself
            return True
        elif self.during_timespan(other.start_date - timedelta(days=1)):
            return True
        else:
            return False

    def order_dates(self):
        if self.start_date > self.end_date:
            # Swaps the two values around so the start date is always at the beginning
            self.start_date, self.end_date = self.end_date, self.start_date


def merge_timespan(first_time_span, other_time_span):
    # Merges two TimeSpan Objects into a new one does not delete the old ones
    if first_time_span.check_overlap(other_time_span):  # Checks for shared days so the TimeSpan can be merged
        new_start = min(first_time_span.start_date, other_time_span.start_date)  # the start of the new merge
        new_end = max (first_time_span.end_date, other_time_span.end_date)
        return TimeSpan(new_start, new_end)
    else:
        print("Cannot merge the time span two disjoined events")

# TODO remove all the old unused bits of code in this file and rename it


def create_church_servers(filepath, sheet, server_list):
    columns = ["Vorname", "Nachname", "Kürzel", "Schuljahr"]
    table_server = pandas.read_excel(filepath, sheet_name=sheet, header=0, engine="openpyxl",
                                     usecols=columns)
    for x in range(0, table_server.shape[0]):
        server_list.append(ChurchServers(lastname=table_server.at[x, "Nachname"],
                                         firstname=table_server.at[x, "Vorname"],
                                         abbreviation=table_server.at[x, "Kürzel"]))


def create_availability(church_server):
    # Placeholder to create a state to check if someone is available
    # currently disabled to reenable change first statement after if to False
    availability_list = [date(2021, 11, 1).isoformat(), date(2021, 12, 1).isoformat(), date(2021, 11, 1).isoformat(),
                         date(2021, 11, 1).isoformat()]
    if random.randint(0, 2) == 1:
        church_server.unavailable.append(availability_list[random.randint(0, 3)])
    else:
        pass


def is_available(church_server, church_service):
    if church_service.date.isoformat() in church_server.unavailable:
        church_server.is_available = False
    else:
        church_server.is_available = True


def is_assigned(check_church_service, church_server):
    # the Function checks if a ChurchServer is already assigned to the Service
    if check_church_service.ListServingMD.count(church_server.abbreviation) == 0:
        church_server.is_allocated = False
    else:
        church_server.is_allocated = True


def allocation_md(current_church_service, List = List_MD):
    while len(current_church_service.ListServingMD) < current_church_service.count:
        selected_md = random.choice(List_MD)
        is_available(selected_md, current_church_service)
        if selected_md.is_available:
            is_assigned(current_church_service, selected_md)
            if not selected_md.is_allocated:
                current_church_service.ListServingMD.append(selected_md.abbreviation)
                selected_md.counter = selected_md.counter + 1
            else:
                pass
        else:
            pass
    else:
        pass



# Functions needed to output the Services to a Docx file


def print_plan_docx(list_services, document_name):
    document = Document()
    document.add_heading("Messdienerplan", 1)
    length = len(List_Services)
    for service_number in range(0, length):
        current_service = list_services[service_number]
        document.add_paragraph(str(current_service.date))
        new_string = "Messe: "
        # puts a string together in order to avoid a newline for every church server
        for x in range(0, current_service.count):
            new_string = new_string + str(current_service.ListServingMD[x]) + ", "
        document.add_paragraph(new_string)
    document.save(document_name)


def statistical_analysis():
    print(List_Services[1])
    print(List_MD[1].unavailable)
    print("mean", mean(statistic_list))
    print("median", median(statistic_list))
    print("standard deviation", stdev(statistic_list))
    print("max", max(statistic_list))
    print("min", min(statistic_list))


# Definition of lists and variables needed later
if __name__ == "__main__":
    church_service_amount = 20
    church_servers_amount = 100

    # Generation of ChurchServers
    create_church_servers("Liste_Messdiener_Computer.xlsx", "Kinder", List_MD)
    create_church_servers("Liste_Messdiener_Computer.xlsx", "Leiter", List_MD)

    # Generates church_service_amount of church services
    for number_services in range(0, church_service_amount):
        List_Services.append(ChurchService(8, datetime(2021, 11, 1, 18, 30)))

    # Allocation of Church Server to Church Services
    for selected_church_server in range(0, church_service_amount):
        allocation_md(List_Services[selected_church_server])
        # print(List_Services[selected_church_server].ListServingMD)


    # Adds the counter values of each Church Server to a list
    for selected_church_server_stats in range(0, church_servers_amount):
        # print(List_MD[selected_church_server_stats].name, List_MD[selected_church_server_stats].counter)
        statistic_list.append(List_MD[selected_church_server_stats].counter)

    # statistic Analysis of the allocation of Church Servants


    #print_plan_docx(List_Services, "test.docx")


    a = datetime(2022,5,1).date()
    b = datetime(2022,5,2).date()
    c = datetime(2022,5,3).date()
    d = datetime(2022,5,4).date()
    af = TimeSpan(a,b)
    bf = TimeSpan(c,c)
    cf = TimeSpan(b,d)
    df = TimeSpan(a,d)
    print(df.description)
    print(merge_timespan(af,df).description)
    print(af)
    print(df)
