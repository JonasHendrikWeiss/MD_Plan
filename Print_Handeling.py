from Zuordnung import ChurchService
from Storage_Operations import data_storage, unpickle_storage
from docx import Document


def print_to_docx(list_to_print, path_template ="Template_plan.docx", name_output="plan.docx"):
    document = Document(str(path_template))
    if len(document.tables) == 1:  # checks if there is only one table else it generates a new table
        table = document.tables[0]
    else:
        table = document.add_table(1, 3)

    amount_services = len(list_to_print)


    for number in range(amount_services):
        if number != 0:  # ensures that the table doesnt run out of rows, by adding one every time after the first add
            table.add_row()
        service = list_to_print[number] # selects each service one time

        table.cell(number, 0).text = service.date.strftime("%d.%m.%Y")  # Formatting the time objects in nice strings
        table.cell(number, 1).text = service.time.strftime('%H:%M')

        text_abbreviations = " "  # A text to which all abbreviations of all MDs are added
        for amount in range(len(service.current_churchservers)):
            text_abbreviations = text_abbreviations + service.current_churchservers[amount].abbreviation

            if amount == len(service.current_churchservers)-1:  # handles the seperation of the Words
                text_abbreviations = text_abbreviations + "\n"
            else:
                text_abbreviations = text_abbreviations + ", "
            table.cell(number, 3).text = text_abbreviations

    # Saves the document with the given name
    document.save(name_output)

teststorage = unpickle_storage()
print_to_docx(teststorage.list_services)
